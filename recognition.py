from pdf2docx import parse
import os
import shutil
import subprocess
from docx import Document

def convert_to_docx(path, save_path):
#     ## 原地址
#     path = r"/home/albay/djs/data/Purchase/2023/05/10/"
#     ## 存储地址
#     save_path = r"/home/albay/djs/data/Purchase/2023/05/10/docx/"
    
    if not os.path.exists(save_path):
        os.mkdir(save_path)
    
    for file in os.listdir(path):
        ## 后缀
        suff = os.path.splitext(file)[1]
        
        ## pdf转docx
        if suff == '.pdf': 
            file_name = os.path.splitext(file)[0]
            pdf_file = path + file_name + suff
            docx_file = save_path + file_name + '.docx'
            parse(pdf_file, docx_file)
        ## doc转docx
        elif suff == '.doc': 
            file_name = os.path.splitext(file)[0]
            doc_file = path + file_name + suff
            os.system("soffice --headless --convert-to docx %s --outdir %s" % (doc_file, save_path))
        ## docx复制到储存地址
        elif suff == '.docx':
            copy_file = path + file
            shutil.copy(copy_file, save_path)
        else: 
            continue
        
def recognition(file):    

    #check = ['招标文件', '磋商文件', '项目编号', '项目名称'， '目录'，'文件格式']

    check = '文件格式'
    for para in file.paragraphs: 
        if check in para.text: 
            return True   
    return False
    
#     for i in range(len(file.paragraphs)):
#         print(str(i), file.paragraphs[i].text)
#     for file in os.listdir(path):
#         read_file = path + file
        
def recogAll(path, save_path, trash_path):
    ## 读取地址
#     path = r"/home/albay/djs/data/Purchase/2022/12/01/docx/"
#     save_path = r"/home/albay/djs/data/Purchase_2/2022/12/01/"
#     trash_path = r"/home/albay/djs/trash/2022/12/01/"
    
    if not os.path.exists(save_path):
        os.mkdir(save_path)
        
    if not os.path.exists(trash_path):
        os.mkdir(trash_path)
    
    for file in os.listdir(path): 
        try: 
            read_file = Document(path + file)
            tf = recognition(read_file)
            if tf == True: 
                copy_file = path + file
                shutil.copy(copy_file, save_path)
                print("{0} is copied and saved".format(copy_file))
            else:
                copy_file = path + file
                shutil.copy(copy_file, trash_path)
        except Exception as e:
            print('Error: ', e)

    
def main():
    date = "\\18\\"
    path1 = r"C:\Users\Administrator\Desktop"
    
    read_path = path1 + date
    save_path =  read_path + "\\docx\\"
    
    path2 = r"C:\Users\Administrator\Desktop\reg"
    bid_path = path2 + date
    trash_path = r"C:\Users\Administrator\Desktop\trash" + date
    
    print(read_path)
    print(save_path)
    print(bid_path)
    print(trash_path)
    
    convert_to_docx(read_path, save_path)
    recogAll(save_path, bid_path, trash_path)

if __name__ == '__main__':
    main()