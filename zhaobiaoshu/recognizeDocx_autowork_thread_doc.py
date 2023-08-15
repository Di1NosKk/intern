import os
import shutil
import subprocess
from docx import Document
import time
import datetime
import signal
from contextlib import contextmanager
import pytz
import aspose.words as aw
from concurrent.futures import ProcessPoolExecutor


class TimeoutException(Exception): pass

@contextmanager
def time_limit(seconds):
    def signal_handler(signum, frame):
        raise TimeoutException("Timed out!")
    signal.signal(signal.SIGALRM, signal_handler)
    signal.alarm(seconds)
    try:
        yield
    finally:
        signal.alarm(0)
    
    
def docxtodocx(copy_file, save_path, file):
    print("当前docx：{}----> started at {}".format(file, time.strftime('%Y-%m-%d %H:%M:%S', time.localtime())))
    shutil.copy(copy_file, save_path)
    recogAll(save_path, bid_path, file)

def convert_to_docx(path, save_path):
    if not os.path.exists(save_path):
        os.makedirs(save_path, exist_ok=True)
    
    for file in os.listdir(path):
        ## 后缀
        suff = os.path.splitext(file)[1]
            
        # doc转docx
        if suff == '.doc': 
            file_name = os.path.splitext(file)[0]
            doc_file = path + file_name + suff
            docx_file = save_path + file_name + '.docx'
            print("当前doc：{}----> started at {}".format(file, time.strftime('%Y-%m-%d %H:%M:%S', time.localtime())))
            try:
                with time_limit(30):
                    doc = aw.Document(doc_file)
                    doc.save(docx_file)
            except Exception as e:
                print("Time out!{} and file is {}".format(e, file))
            recogAll(save_path, bid_path, file_name + ".docx")
                
        ## docx复制到储存地址
        elif suff == '.docx':
            copy_file = path + file
            docxtodocx(copy_file, save_path, file)

        
def recognition(file):    

    #check = ['招标文件', '磋商文件', '项目编号', '项目名称'， '目录'，'文件格式']

    check = '文件格式'
    for para in file.paragraphs: 
        if check in para.text: 
            return True   
    return False
    
        
def recogAll(path, save_path, file):
    ## 读取地址
    
    if not os.path.exists(save_path):
        os.makedirs(save_path, exist_ok=True)

    try: 
        read_file = Document(path + file)
        tf = recognition(read_file)
        if tf == True: 
            copy_file = path + file
            shutil.copy(copy_file, save_path)
            print("{0} is copied and saved in {1}".format(file, save_path))
    except Exception as e:
        print('Error: ', e)

    
def main():
    convert_to_docx(read_path, save_path)
    print("finished")

if __name__ == '__main__':
    today = datetime.datetime.today().astimezone(pytz.timezone('Etc/GMT-8'))
    oneday = datetime.timedelta(days=1)
    yesterday = (today - oneday).strftime('%Y/%m/%d')
    lst = yesterday.split("/")
    year = lst[0]
    month = lst[1]
    day = lst[2]
    date = yesterday + "/"
    path1 = r"/data/data/purchase_resource/Purchase/"
    
    read_path = path1 + date
    save_path = r"/data/data/zhaobiao_book_tmp_doc/" + date
    
    path2 = r"/data/data/purchase_resource/Purchase_zhaobiao_book/"
    bid_path = path2 + date
    
    print(read_path)
    print(save_path)
    print(bid_path)
    
    main()

